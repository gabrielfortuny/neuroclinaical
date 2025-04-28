import os
import fitz
import traceback
import zipfile
from docx import Document
from flask import current_app
#from app.__init__ import db
from app import db
from app.models import Report, ExtractedImage
from app.services.data_upload.nlpRequestHandler import handle_drugadmin_request, handle_seizure_request, handle_summary_request
from app.services.data_upload.uploadUtilities import (
    extract_days_from_text,
    store_drugs_array,
    store_seizures_array,
)

# INDIVIDUAL UPLOAD HANDLERS: DON'T USE DIRECTLY


class pdf_upload_handler:
    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            text = []
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text.append(page.get_text())
            return "\n".join(text)
        except Exception as e:
            current_app.logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"PDF extraction failed: {str(e)}")

    def __call__(self, file_path: str, storage_path: str, report_id: int) -> str:
        """Extract text from a PDF file."""
        current_app.logger.info(f"Extracting text from PDF: {file_path}")
        text = self.extract_text_from_pdf(file_path)
        self.extract_image_from_pdf(file_path, storage_path, report_id)
        current_app.logger.info(f"Extracted {len(text)} characters from PDF")
        return text

    def extract_image_from_pdf(self, filepath:str, storage_path:str, report_id: int):
        try:
            doc = fitz.open(filepath)
            image_num = 0
            for page in range(len(doc)):
                images = doc[page].get_images()
                for image in images:
                    im = ExtractedImage(report_id=report_id)
                    db.session.add(im)
                    xref = image[0] 
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    ext = base_image["ext"]
                    final_path = f"{storage_path}/{report_id}_image{image_num}.{ext}"
                    im.file_path = final_path
                    image_num += 1
                    with open(final_path, "wb") as f:
                        f.write(image_bytes)
            db.session.commit()
        except Exception as e:
            current_app.logger.error(f"Error extracting images from PDF: {str(e)}")
            raise Exception(f"PDF extraction failed: {str(e)}")
        
class docx_upload_handler:

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            extracted_text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    extracted_text.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        extracted_text.append(row_text)

            return "\n".join(extracted_text)
        except Exception as e:
            current_app.logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"DOCX extraction failed: {str(e)}")

    def extract_image_from_docx(self, filepath:str, storage_path:str, report_id: int):
        try:
            with zipfile.ZipFile(filepath, 'r') as docx:
                image_num = 0
                for item in docx.namelist():
                    if item.startswith("word/media/"):
                        image = ExtractedImage(report_id=report_id)
                        db.session.add(image)
                        #If failing: Commit to db needed?
                        image_data = docx.read(item)
                        ext = os.path.splitext(item)[1]
                        final_path = f"{storage_path}/{report_id}_image{image_num}.{ext}"
                        image.file_path = final_path
                        image_num += 1
                        with open(final_path, 'wb') as f:
                            f.write(image_data)
                db.session.commit()
                        
        except Exception as e:
            current_app.logger.error(f"Error extracting images from DOCX: {str(e)}")
            raise Exception(f"DOCX extraction failed: {str(e)}")


    def __call__(self, file_path: str, storage_path: str, report_id: int) -> str:
        """Extract text from a DOCX file."""
        current_app.logger.info(f"Extracting text from DOCX: {file_path}")
        text = self.extract_text_from_docx(file_path)
        self.extract_image_from_docx(file_path, storage_path, report_id)
        current_app.logger.info(f"Extracted {len(text)} characters from DOCX")
        return text


supported_file_types = {
    "pdf": pdf_upload_handler(),
    "docx": docx_upload_handler(),
}

# DIRECTLY USABLE UPLOAD HANDLERS:


def upload_controller(content_ext: str,
                      file_path: str,
                      p_id: int,
                      report: Report) -> bool:
    """
    • extracts text (pdf / docx)
    • asks Ollama for a summary
    • writes the summary back to the Report row
    """
    current_app.logger.info(f"Starting upload_controller for {content_ext}: {file_path}")

    # ---- 1.  scrape --------------------------------------------------------
    storage_path = os.path.dirname(file_path)          # …reports/<patient_id>
    extractor     = supported_file_types.get(content_ext)
    if extractor is None:
        current_app.logger.warning(f"Unsupported file type: {content_ext}")
        return False

    try:
        # NB: docx handler needs storage_path & report.id to pull out images
        extracted_text = extractor(file_path, storage_path, report.id)
    except Exception:
        current_app.logger.error("Text extraction failed", exc_info=True)
        return False

    if not extracted_text:
        current_app.logger.warning("No text extracted")
        return False

    # ---- 2.  summarise -----------------------------------------------------
    summary = handle_summary_request(extracted_text)
    if not summary:                                   # model may time‑out
        summary = extracted_text[:400] + "…"          # cheap fall‑back

    # ---- 3.  persist -------------------------------------------------------
    try:
        report.summary = summary.strip()
        db.session.commit()
        days = extract_days_from_text(extracted_text)
        current_app.logger.info(f"Checking day error {days}")
        
        seizures = handle_seizure_request(days)
        store_seizures_array(seizures, p_id)
        drug_admin = handle_drugadmin_request(days)
        store_drugs_array(drug_admin, p_id)
        db.session.commit()
        current_app.logger.info("Report summary saved")
        return True

    except Exception:
        db.session.rollback()
        current_app.logger.error("DB commit failed", exc_info=True)
        return False
