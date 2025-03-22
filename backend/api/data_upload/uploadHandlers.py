import os
import fitz
import typing
import zipfile
import base64
from docx import Document
from io import BytesIO
from flask import current_app
import sqlalchemy
from nlpRequestHandler import (
    handle_seizure_request,
    handle_summary_request,
    handle_drugadmin_request,
)
from uploadUtilities import (
    extract_days_from_text,
    store_drugs_array,
    store_seizures_array,
)
from app.__init__ import db
from app.models import Report

# INDIVIDUAL UPLOAD HANDLERS: DON'T USE DIRECTLY


class pdf_upload_handler:
    def extract_text_from_pdf(self, file_path: str) -> str:
        text = []
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text.append(page.get_text())
        return "\n".join(text)

    def __call__(self, data: str, p_id: int) -> bool:
        # TODO: Write Docstring
        """
        Description:

        Requires:
        Query is a string for the

        Modifies:

        Effects:

        @param data: String to be used for pdf
        """
        try:
            pdf_base64 = data
            pdf_binary = base64.b64decode(pdf_base64)  # Decode Base64 Encoding
            report = Report(
                patient_id=p_id
            )  # Create report object to extract ID in order to create a folder with same name
            db.session.add(report)
            db.session.commit()
            os.makedirs(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}", exist_ok=True
            )  # Create file path (It's ok if this already exists for a patient)
            os.makedirs(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}",
                exist_ok=False,
            )  # Create folder with report ID (Should never exist)
            # Create Report file in aformentioned folder
            with open(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}/report.pdf",
                "wb",
            ) as f:
                f.write(pdf_binary)
            report.filetype = "pdf"
            report.filepath = f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}"  # Update report filepath attribute to fit new folder created

            text = self.extract_text_from_pdf(
                report.filepath
            )  # Extract raw text from the docx

            with open(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}/report.txt",
                "w",
                encoding="utf-8",
            ) as f:
                f.write(text)
            # Store raw text is same folder as report.txt for easy universal access
            days_dict = extract_days_from_text(
                text
            )  # Extract all days from text into a dictionary
            count = 0

            # TODO: EXTRACT IMAGES

            for day, content in days_dict.items():  # Store all days as files

                day_number = int((day.split(" "))[1])

                filename = f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}/{day}.txt"
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(content)
            # HANDLE SEIZURE
            seizures = handle_seizure_request(days_dict)
            if not store_seizures_array(seizures, p_id):
                raise sqlalchemy.exc.SQLAlchemyError

            # HANDLE DRUGS
            drugs = handle_drugadmin_request(days_dict)
            if not store_drugs_array(drugs, p_id):
                raise sqlalchemy.exc.SQLAlchemyError

            # HANDLE SUMMARY
            report.days = count
            report.summary = handle_summary_request(text)
            db.session.commit()  # Save changes to out DB "report" object
            return True
        except Exception as err:
            return False


class docx_upload_handler:
    def extract_text_from_docx(self, file_path: str) -> str:
        # TODO: Write Docstring
        """
        Description:

        Requires:
        Query is a string for the

        Modifies:

        Effects:

        @param data: String to be used for pdf
        """
        doc = Document(file_path)
        extracted_text = []

        # Extract paragraphs (handling section headers)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                extracted_text.append(text)

        # Extract tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    extracted_text.append(row_text)

        return "\n".join(extracted_text)

    def __call__(self, data: str, p_id: int) -> bool:
        """
        Description:

        Requires:

        Modifies:

        Effects:

        @param data: String to be used for docx file
        @param p_id: The Patient ID this report belongs to
        @return: bool indicating successful execution
        """
        try:
            docx_base64 = data
            docx_binary = base64.b64decode(docx_base64)  # Decode Base64 Encoding
            report = Report(
                patient_id=p_id
            )  # Create report object to extract ID in order to create a folder with same name
            db.session.add(report)
            db.session.commit()
            os.makedirs(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}", exist_ok=True
            )  # Create file path (It's ok if this already exists for a patient)
            os.makedirs(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}",
                exist_ok=False,
            )  # Create folder with report ID (Should never exist)
            # Create Report file in aformentioned folder
            with open(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}/report.docx",
                "wb",
            ) as f:
                f.write(docx_binary)
            report.filetype = "docx"
            report.filepath = f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}"  # Update report filepath attribute to fit new folder created

            text = self.extract_text_from_docx(
                report.filepath
            )  # Extract raw text from the docx

            with open(
                f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}/report.txt",
                "w",
                encoding="utf-8",
            ) as f:
                f.write(text)
            # Store raw text is same folder as report.txt for easy universal access
            days_dict = extract_days_from_text(
                text
            )  # Extract all days from text into a dictionary
            count = 0

            # TODO: EXTRACT IMAGES

            for day, content in days_dict.items():  # Store all days as files

                day_number = int((day.split(" "))[1])

                filename = f"{current_app.config["UPLOAD_FOLDER"]}/{p_id}/{report.id}/{day}.txt"
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(content)
            # HANDLE SEIZURE
            seizures = handle_seizure_request(days_dict)
            if not store_seizures_array(seizures, p_id):
                raise sqlalchemy.exc.SQLAlchemyError

            # HANDLE DRUGS
            drugs = handle_drugadmin_request(days_dict)
            if not store_drugs_array(drugs, p_id):
                raise sqlalchemy.exc.SQLAlchemyError

            # HANDLE SUMMARY
            report.days = count
            report.summary = handle_summary_request(text)
            db.session.commit()  # Save changes to out DB "report" object
            return True
        except Exception as err:
            return False


supported_file_types = {
    "application/pdf": pdf_upload_handler(),
    "application/docx": docx_upload_handler(),
}

# DIRECTLY USABLE UPLOAD HANDLERS:


def upload_controller(content_type: str, content: str, patient_id: int) -> bool:
    """
    Description: Preforms all steps needed to upload and analyze a document

    Requires:
    Query is a string for the

    Modifies:

    Effects:

    @param content_type: String containing the content type from the request header
    @param content: String containing the file data in base 64 encoding
    @param patient_id: The patient who this report will belong to
    """

    if content_type in supported_file_types:
        supported_file_types[content_type](content, patient_id)
        return True
    return False
